import os
import io
import json
import re
import uuid
import datetime
import zipfile
import math
from flask import Flask, request, render_template, send_file, flash, redirect, url_for, jsonify
from werkzeug.utils import secure_filename
from groq import Groq
import pdfplumber
import docx
from docx import Document
from PIL import Image, ImageDraw, ImageFont
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    import mammoth
except ImportError:
    mammoth = None

try:
    import pythoncom
except ImportError:
    pythoncom = None

app = Flask(__name__)
app.secret_key = "super_secret_ablbl_key_prototype3"

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

# Load local secure .env file if present
if os.path.exists(".env"):
    try:
        with open(".env", "r") as env_file:
            for line in env_file:
                if "=" in line and not line.strip().startswith("#"):
                    key, val = line.strip().split("=", 1)
                    os.environ[key.strip()] = val.strip()
    except Exception:
        pass

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")




LOGO_ABLBL = os.path.join("static", "logo.png")
ICON_WEB = os.path.join("static", "website.png")
ICON_LIN = os.path.join("static", "linkedin.png")
ICON_INST = os.path.join("static", "instagram.png")
ICON_DIR = os.path.join("static", "directory.png")
ICON_BRIEFCASE = os.path.join("static", "briefcase.png")
ICON_GEAR = os.path.join("static", "gear.png")
ICON_PIN = os.path.join("static", "pin.png")
ICON_DIAMOND = os.path.join("static", "diamond.png")

def generate_all_icons():
    os.makedirs("static", exist_ok=True)
    size = (128, 128)
    color_blue = (69, 129, 181, 255) # #4581B5
    
    # 1. Briefcase icon
    if not os.path.exists(ICON_BRIEFCASE):
        img = Image.new("RGBA", size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle([24, 44, 104, 104], radius=10, outline=color_blue, width=6)
        draw.rounded_rectangle([48, 24, 80, 44], radius=6, outline=color_blue, width=6)
        draw.rectangle([60, 68, 68, 80], fill=color_blue)
        img.save(ICON_BRIEFCASE)
        
    # 2. Gear icon
    if not os.path.exists(ICON_GEAR):
        img = Image.new("RGBA", size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.ellipse([44, 44, 84, 84], outline=color_blue, width=6)
        draw.ellipse([58, 58, 70, 70], outline=color_blue, width=4)
        for i in range(8):
            angle = i * (math.pi / 4)
            cx = 64 + 48 * math.cos(angle)
            cy = 64 + 48 * math.sin(angle)
            draw.ellipse([cx-8, cy-8, cx+8, cy+8], fill=color_blue)
        img.save(ICON_GEAR)
        
    # 3. Location pin icon
    if not os.path.exists(ICON_PIN):
        img = Image.new("RGBA", size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.ellipse([34, 24, 94, 84], outline=color_blue, width=6)
        draw.ellipse([54, 44, 74, 64], outline=color_blue, width=6)
        draw.polygon([(40, 74), (64, 114), (88, 74)], fill=color_blue)
        img.save(ICON_PIN)
        
    # 4. Diamond icon
    if not os.path.exists(ICON_DIAMOND):
        img = Image.new("RGBA", size, (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.polygon([(64, 24), (104, 64), (64, 104), (24, 64)], outline=color_blue, width=6)
        img.save(ICON_DIAMOND)
        
    # Footer white transparent icons
    size_f = (64, 64)
    color_white = (255, 255, 255, 255)
    color_navy = (13, 25, 41, 255)
    
    # 5. Website
    if not os.path.exists(ICON_WEB):
        img = Image.new("RGBA", size_f, (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.ellipse([8, 8, 56, 56], outline=color_white, width=4)
        draw.line([8, 32, 56, 32], fill=color_white, width=4)
        draw.line([32, 8, 32, 56], fill=color_white, width=4)
        draw.ellipse([18, 8, 46, 56], outline=color_white, width=3)
        img.save(ICON_WEB)
        
    # 6. LinkedIn
    if not os.path.exists(ICON_LIN):
        img = Image.new("RGBA", size_f, (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle([8, 8, 56, 56], radius=8, fill=color_white)
        try: font = ImageFont.truetype("arial.ttf", 36)
        except: font = ImageFont.load_default()
        draw.text((15, 8), "in", fill=color_navy, font=font)
        img.save(ICON_LIN)
        
    # 7. Instagram
    if not os.path.exists(ICON_INST):
        img = Image.new("RGBA", size_f, (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle([8, 8, 56, 56], radius=12, outline=color_white, width=4)
        draw.ellipse([20, 20, 44, 44], outline=color_white, width=4)
        draw.ellipse([42, 18, 46, 22], fill=color_white)
        img.save(ICON_INST)
        
    # 8. Directory
    if not os.path.exists(ICON_DIR):
        img = Image.new("RGBA", size_f, (255, 255, 255, 0))
        draw = ImageDraw.Draw(img)
        draw.rounded_rectangle([8, 8, 56, 56], radius=8, outline=color_white, width=4)
        draw.line([16, 20, 48, 20], fill=color_white, width=3)
        draw.line([16, 32, 48, 32], fill=color_white, width=3)
        draw.line([16, 44, 36, 44], fill=color_white, width=3)
        img.save(ICON_DIR)

generate_all_icons()

AI_PROMPT = """
You are a world-class corporate HR copywriter for Aditya Birla Lifestyle Brands Limited (ABLBL).
Transform the legacy Job Description into structured JSON.
Return ONLY valid JSON — no markdown, no explanation.

LEGACY JD:
{legacy_text}

Return exactly this JSON structure:
{{
  "job_title": "",
  "department": "",
  "location": "",
  "job_purpose": "",
  "skills_stack": ["Skill 1", "Skill 2", "Skill 3", "Skill 4", "Skill 5", "Skill 6"],
  "why_exciting": [
    "Bullet 1 on why this role is exciting", 
    "Bullet 2 on why this role is exciting", 
    "Bullet 3 on why this role is exciting"
  ],
  "kra_list": [
    {{"heading": "KRA 1", "bullets": ["Action 1", "Action 2"]}},
    {{"heading": "KRA 2", "bullets": ["Action 1", "Action 2"]}}
  ]
}}

STRICT RULES:
1. "job_title": Extract from "Designation of the Employee".
2. "department": Extract from "Department".
3. "location": Extract from "Location".
4. "job_purpose": Extract from "Job Purpose".
5. "kra_list": Extract from "Key Result Areas & Supporting Actions". Create 3 to 6 KRAs based on the legacy text. Each KRA MUST have exactly 2 bullet points under it summarizing the supporting actions.
6. "skills_stack": Infer exactly 6 key skills relevant to the role.
7. "why_exciting": Infer exactly 3 compelling bullet points explaining why this role is exciting and strategic.
8. BRAND SAFETY: Under no circumstances mention "Pivot", "Birla Pivot", or "Aditya Birla Pivot" anywhere in the JSON output! Instead, write "ABLBL" or "Aditya Birla Lifestyle Brands Limited" (ABLBL).
"""

def clean_text(text):
    if not isinstance(text, str):
        return text
    # Proactively replace Pivot terms with ABLBL
    text = re.sub(r'\bBirla\s+Pivot\b', 'ABLBL', text, flags=re.IGNORECASE)
    text = re.sub(r'\bAditya\s+Birla\s+Pivot\b', 'ABLBL', text, flags=re.IGNORECASE)
    text = re.sub(r'\bPivot\b', 'ABLBL', text, flags=re.IGNORECASE)
    return text

def clean_data_recursively(data):
    if isinstance(data, dict):
        return {k: clean_data_recursively(v) for k, v in data.items()}
    elif isinstance(data, list):
        return [clean_data_recursively(x) for x in data]
    elif isinstance(data, str):
        return clean_text(data)
    return data

def extract_text_from_bytes(file_bytes, filename):
    ext = filename.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        txt = ""
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for p in pdf.pages: txt += p.extract_text() or ""
        return txt
    if ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    if ext == "docx":
        if mammoth:
            try:
                result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
                if result.value.strip(): return result.value
            except Exception: pass
        try:
            doc = Document(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception: pass
    if ext == "doc":
        if mammoth:
            try:
                result = mammoth.extract_raw_text(io.BytesIO(file_bytes))
                if result.value.strip(): return result.value
            except Exception: pass
        try:
            chunks = re.findall(rb'[\x20-\x7E\n\r\t]{4,}', file_bytes)
            text = "\n".join(c.decode("ascii", errors="ignore") for c in chunks)
            lines = [l.strip() for l in text.splitlines() if len(l.strip()) >= 3 and any(c.isalnum() for c in l)]
            if lines: return "\n".join(lines)

        except Exception: pass
    return ""

def extract_json(raw):
    raw = re.sub(r"^```(?:json)?\s*", "", raw.strip(), flags=re.DOTALL)
    raw = re.sub(r"\s*```\s*$", "", raw.strip(), flags=re.DOTALL)
    try: return json.loads(raw)
    except json.JSONDecodeError: pass
    for s_char, e_char in [("{", "}"), ("[", "]")]:
        start = raw.find(s_char)
        end = raw.rfind(e_char)
        if start != -1 and end != -1 and end > start:
            try: return json.loads(raw[start:end+1])
            except json.JSONDecodeError: pass
    raise json.JSONDecodeError("No valid JSON found")

def call_groq(text_payload, api_key=None):
    use_key = api_key if api_key else GROQ_API_KEY
    client = Groq(api_key=use_key)
    resp = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": AI_PROMPT.format(legacy_text=text_payload[:8000])}],
        temperature=0.2,
        max_tokens=4096,
    )
    raw = resp.choices[0].message.content.strip()
    try:
        return extract_json(raw)
    except json.JSONDecodeError:
        strict = f"Output ONLY a JSON object matching the requested schema. No markdown. Extract from: {text_payload[:4000]}"
        resp2 = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": strict}],
            temperature=0.1,
            max_tokens=3000,
        )
        return extract_json(resp2.choices[0].message.content.strip())

def summarize_preferred_qualification(text, api_key=None):
    if not text.strip():
        return []
    use_key = api_key if api_key else GROQ_API_KEY
    client = Groq(api_key=use_key)
    prompt = f"""
You are a professional HR assistant for Aditya Birla Lifestyle Brands Limited (ABLBL).
Summarize the following raw text regarding preferred qualifications/certifications/experience for a role into a concise list of professional bullet points (maximum 5 points).
Each point must be a clean requirement (e.g., "MBA from a Tier-1 business school is preferred", "Strong experience in apparel retail operations", etc.).
Return ONLY a valid JSON list of strings. No explanation.

RAW TEXT:
{text}
"""
    try:
        resp = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=1000,
        )
        raw = resp.choices[0].message.content.strip()
        data = extract_json(raw)
        if isinstance(data, list):
            return [str(x).strip() for x in data if str(x).strip()]
        return []
    except Exception:
        # Fallback: parse sentences/lines
        lines = [line.strip().lstrip("*-•❖ ").strip() for line in text.split("\n") if line.strip()]
        return lines[:5]

def rgb(h):
    h = h.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def shd_cell(cell, color):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color.lstrip("#"))
    pr.append(shd)

def set_cell_margins(cell, top=0, bottom=0, start=0, end=0):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('tcMar'):
            tcPr.remove(child)
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', start), ('right', end)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_valign_center(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)

def clear_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            tblBorders.append(border)
        tblPr[0].append(tblBorders)

def add_image_hyperlink(paragraph, image_path, url, width=Cm(0.6)):
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    
    temp_run = paragraph.add_run()
    temp_run.add_picture(image_path, width=width)
    drawing = temp_run._r.find(qn('w:drawing'))
    
    if drawing is not None:
        temp_run._r.remove(drawing)
        new_run.append(drawing)
        
    p_element = paragraph._element
    p_element.remove(temp_run._r)
    
    hyperlink.append(new_run)
    p_element.append(hyperlink)

def create_first_page_header(header_obj):
    header_obj.is_linked_to_previous = False
    table = header_obj.add_table(rows=1, cols=1, width=Inches(8.5))
    clear_table_borders(table)
    table.columns[0].width = Inches(8.5)
    
    cell = table.cell(0, 0)
    cell.width = Inches(8.5)
    shd_cell(cell, "0D1929")
    set_cell_margins(cell, top=360, bottom=360, start=0, end=0)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.85)
    p.paragraph_format.right_indent = Inches(0.85)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_ABLBL):
        p.add_run().add_picture(LOGO_ABLBL, width=Inches(1.2))
        
    p_title = cell.add_paragraph()
    p_title.paragraph_format.space_before = Pt(20)
    p_title.paragraph_format.space_after = Pt(0)
    p_title.paragraph_format.left_indent = Inches(0.85)
    p_title.paragraph_format.right_indent = Inches(0.85)
    r_title = p_title.add_run("ROLE PLAYBOOK")
    r_title.bold = True
    r_title.font.size = Pt(28)
    r_title.font.color.rgb = rgb("FFFFFF")
    r_title.font.name = "Calibri"

def create_subsequent_page_header(header_obj, job_title):
    header_obj.is_linked_to_previous = False
    table = header_obj.add_table(rows=1, cols=2, width=Inches(8.5))
    clear_table_borders(table)
    
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(6.0)
    
    cell_logo = table.cell(0, 0)
    cell_logo.width = Inches(2.5)
    shd_cell(cell_logo, "0D1929")
    set_cell_margins(cell_logo, top=180, bottom=180, start=0, end=0)
    
    p_logo = cell_logo.paragraphs[0]
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after = Pt(0)
    p_logo.paragraph_format.left_indent = Inches(0.85)
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(LOGO_ABLBL):
        p_logo.add_run().add_picture(LOGO_ABLBL, width=Inches(0.9))
        
    cell_title = table.cell(0, 1)
    cell_title.width = Inches(6.0)
    shd_cell(cell_title, "0D1929")
    set_cell_margins(cell_title, top=180, bottom=180, start=0, end=0)
    
    p_title = cell_title.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(0)
    p_title.paragraph_format.right_indent = Inches(0.85)
    p_title.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_title = p_title.add_run(f"Role Playbook — {job_title}")
    r_title.bold = True
    r_title.font.size = Pt(12)
    r_title.font.color.rgb = rgb("FFFFFF")
    r_title.font.name = "Calibri"

def create_footer(footer_obj):
    footer_obj.is_linked_to_previous = False
    table = footer_obj.add_table(rows=1, cols=1, width=Inches(8.5))
    clear_table_borders(table)
    table.columns[0].width = Inches(8.5)
    
    cell = table.cell(0, 0)
    cell.width = Inches(8.5)
    shd_cell(cell, "0D1929")
    set_cell_margins(cell, top=180, bottom=180, start=0, end=0)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.85)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    icon_w = Inches(0.24)
    if os.path.exists(ICON_WEB):
        add_image_hyperlink(p, ICON_WEB, "https://www.ablbl.in/", width=icon_w)
    
    p.add_run("      ")
    
    if os.path.exists(ICON_LIN):
        add_image_hyperlink(p, ICON_LIN, "https://www.linkedin.com/company/aditya-birla-lifestyle-brands-limited/posts/?feedView=all", width=icon_w)





def add_body_paragraph(doc, text="", space_before=Pt(0), space_after=Pt(3), line_spacing=1.15):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.75)
    p.paragraph_format.right_indent = Inches(0.75)
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing = line_spacing
    if text:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(10)
        r.font.color.rgb = rgb("333333")
    return p

def add_heading_with_bar(doc, title, space_before=Pt(8), space_after=Pt(3)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.75)
    p.paragraph_format.right_indent = Inches(0.75)
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), '1F6CA6')
    pBdr.append(left)
    pPr.append(pBdr)
    
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = rgb("13375B")
    r.font.name = "Calibri"
    return p

def add_bullet_paragraph(doc, text, sym="❖", space_after=Pt(2)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.95)
    p.paragraph_format.right_indent = Inches(0.75)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = space_after
    
    run_sym = p.add_run(f"{sym}   ")
    run_sym.font.name = "Calibri"
    run_sym.font.size = Pt(10)
    run_sym.font.color.rgb = rgb("1F6CA6")
    
    run_text = p.add_run(text)
    run_text.font.name = "Calibri"
    run_text.font.size = Pt(10)
    run_text.font.color.rgb = rgb("333333")
    return p

def add_spacer(doc, height_pt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(height_pt)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    run = p.add_run()
    run.font.size = Pt(1)

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.75)
    p.paragraph_format.right_indent = Inches(0.75)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(1)
    run = p.add_run()
    run.font.size = Pt(1)
    
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'E0E0E0')
    pBdr.append(bottom)
    pPr.append(pBdr)


def build_docx(data, filepath):
    doc = Document()
    
    # 1. Page & Section Margins to 0 for Full-Bleed
    for sec in doc.sections:
        sec.page_width = Inches(8.5)
        sec.page_height = Inches(11.0)
        sec.top_margin = Inches(0)
        sec.bottom_margin = Inches(0)
        sec.left_margin = Inches(0)
        sec.right_margin = Inches(0)
        sec.header_distance = Inches(0)
        sec.footer_distance = Inches(0)
        sec.different_first_page_header_footer = True

    # jt, dept, loc extraction
    jt = data.get("job_title", "Category Lead")
    dept = data.get("department", "Category")
    loc = data.get("location", "Bangalore")

    # 2. Setup Headers and Footers natively
    create_first_page_header(doc.sections[0].first_page_header)
    create_subsequent_page_header(doc.sections[0].header, jt)
    create_footer(doc.sections[0].first_page_footer)
    create_footer(doc.sections[0].footer)

    # 3. PAGE 1 CONTENT
    # Spacer to clear the first page header (banner height = 1.8 inches)
    add_spacer(doc, 110)




    # Card 1: Job Title (full width)
    t_jt = doc.add_table(rows=1, cols=1)
    t_jt.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(t_jt)
    t_jt.columns[0].width = Inches(7.0)
    
    cell = t_jt.cell(0, 0)
    shd_cell(cell, "F4F7FA")
    set_cell_margins(cell, top=140, bottom=140, start=140, end=140)
    set_cell_valign_center(cell)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    
    if os.path.exists(ICON_BRIEFCASE):
        run_icon = p.add_run()
        run_icon.add_picture(ICON_BRIEFCASE, width=Inches(0.35))
        
    p.add_run(" ")
    
    r_l1 = p.add_run("You will join us as ")
    r_l1.font.size = Pt(9.5)
    r_l1.font.color.rgb = rgb("4581B5")
    r_v1 = p.add_run(jt)
    r_v1.bold = True
    r_v1.font.size = Pt(13)
    r_v1.font.color.rgb = rgb("000000")

    add_spacer(doc, 6)

    # Card 2 & 3: Department and Location side-by-side
    t_dl = doc.add_table(rows=1, cols=3)
    t_dl.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(t_dl)
    t_dl.columns[0].width = Inches(3.4)
    t_dl.columns[1].width = Inches(0.2)
    t_dl.columns[2].width = Inches(3.4)
    
    # Department
    cd = t_dl.cell(0, 0)
    shd_cell(cd, "F4F7FA")
    set_cell_margins(cd, top=120, bottom=120, start=120, end=120)
    set_cell_valign_center(cd)
    
    pd = cd.paragraphs[0]
    pd.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pd.paragraph_format.space_before = Pt(0)
    pd.paragraph_format.space_after = Pt(0)
    
    if os.path.exists(ICON_GEAR):
        run_dicon = pd.add_run()
        run_dicon.add_picture(ICON_GEAR, width=Inches(0.3))
        
    pd.add_run(" ")
    
    r_l2 = pd.add_run("You will work in ")
    r_l2.font.size = Pt(9)
    r_l2.font.color.rgb = rgb("4581B5")
    r_v2 = pd.add_run(dept)
    r_v2.bold = True
    r_v2.font.size = Pt(11)
    r_v2.font.color.rgb = rgb("000000")
    
    # Spacer
    cd_space = t_dl.cell(0, 1)
    set_cell_margins(cd_space, top=0, bottom=0, start=0, end=0)
    
    # Location
    cl = t_dl.cell(0, 2)
    shd_cell(cl, "F4F7FA")
    set_cell_margins(cl, top=120, bottom=120, start=120, end=120)
    set_cell_valign_center(cl)
    
    pl = cl.paragraphs[0]
    pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pl.paragraph_format.space_before = Pt(0)
    pl.paragraph_format.space_after = Pt(0)
    
    if os.path.exists(ICON_PIN):
        run_licon = pl.add_run()
        run_licon.add_picture(ICON_PIN, width=Inches(0.3))
        
    pl.add_run(" ")
    
    r_l3 = pl.add_run("You'll be based out of ")
    r_l3.font.size = Pt(9)
    r_l3.font.color.rgb = rgb("4581B5")
    r_v3 = pl.add_run(loc)
    r_v3.bold = True
    r_v3.font.size = Pt(11)
    r_v3.font.color.rgb = rgb("000000")

    add_horizontal_line(doc)

    # About Us Section
    add_heading_with_bar(doc, "About Us", space_before=Pt(8))
    
    p1 = "Aditya Birla Lifestyle Brands Limited, part of the Aditya Birla Group, is a premier apparel leader with a retail footprint exceeding 4.8 million sq. ft. as of late 2025. Its vast network includes 3,315 exclusive outlets and over 45,000 additional points of sale nationwide."
    add_body_paragraph(doc, p1, space_after=Pt(6))
    
    p2 = "The company’s portfolio features iconic names like Louis Philippe, Van Heusen, Allen Solly, and Peter England, all boasting over 25 years of market leadership. These brands excel in formal, casual, and occasion wear through a commitment to design innovation and operational excellence."
    add_body_paragraph(doc, p2, space_after=Pt(6))
    
    p3 = "Beyond its core labels, ABLBL has successfully scaled Van Heusen Innerwear and maintains strategic partnerships with international brands American Eagle, Reebok, and Simon Carter, further diversifying its reach across the premium denim, athleisure, and menswear segments."
    add_body_paragraph(doc, p3, space_after=Pt(6))



    # 4. PAGE 2 CONTENT
    doc.add_page_break()
    
    # Spacer to clear the subsequent page header (banner height = 1.0 inch)
    add_spacer(doc, 60)



    
    # Section: Job Purpose
    add_heading_with_bar(doc, "Job Purpose", space_before=Pt(6))
    add_body_paragraph(doc, data.get("job_purpose", ""))
    
    # Section: Why is this role exciting?
    add_heading_with_bar(doc, "Why is this role exciting?", space_before=Pt(6))
    add_body_paragraph(doc, f"As a {jt}, you are the growth engine behind Aditya Birla Lifestyle Brands Limited's most strategic business decisions. In this role you get to:")
    
    why_bullets = data.get("why_exciting", [])
    for b in why_bullets:
        add_bullet_paragraph(doc, b, sym="❖")



    # Section: What will you do?
    add_heading_with_bar(doc, "What will you do?", space_before=Pt(6))
    
    kras = data.get("kra_list", [])
    t_kra = doc.add_table(rows=(len(kras) + 1) // 2 * 2 - 1 if len(kras) > 0 else 1, cols=3)
    t_kra.alignment = WD_TABLE_ALIGNMENT.CENTER
    clear_table_borders(t_kra)
    t_kra.columns[0].width = Inches(3.4)
    t_kra.columns[1].width = Inches(0.2)
    t_kra.columns[2].width = Inches(3.4)
    
    for idx, kra in enumerate(kras):
        row_idx = (idx // 2) * 2
        col_idx = 0 if idx % 2 == 0 else 2
        
        cell = t_kra.cell(row_idx, col_idx)
        set_cell_margins(cell, top=60, bottom=60, start=60, end=60)
        
        p_kh = cell.paragraphs[0]
        p_kh.paragraph_format.space_before = Pt(4)
        p_kh.paragraph_format.space_after = Pt(4)
        
        r_ks = p_kh.add_run("❖  ")
        r_ks.bold = True
        r_ks.font.size = Pt(11)
        r_ks.font.color.rgb = rgb("1F6CA6")
        
        r_kt = p_kh.add_run(kra.get("heading", ""))
        r_kt.bold = True
        r_kt.font.size = Pt(11)
        r_kt.font.color.rgb = rgb("1F6CA6")
        
        for b in kra.get("bullets", []):
            p_kb = cell.add_paragraph()
            p_kb.paragraph_format.left_indent = Inches(0.25)
            p_kb.paragraph_format.space_after = Pt(4)
            
            r_kbs = p_kb.add_run("❖   ")
            r_kbs.font.size = Pt(9.5)
            r_kbs.font.color.rgb = rgb("777777")
            
            r_kbt = p_kb.add_run(b)
            r_kbt.font.size = Pt(9.5)
            r_kbt.font.color.rgb = rgb("333333")
            
    # Preferred Qualification Section
    pref_qual_list = data.get("preferred_qualifications", [])
    if pref_qual_list:
        add_spacer(doc, 8)
        add_heading_with_bar(doc, "Preferred Qualification", space_before=Pt(6))
        for point in pref_qual_list:
            add_bullet_paragraph(doc, point, sym="❖")

    doc.save(filepath)
    return jt



@app.route("/")
def index():
    return render_template("index.html")

@app.route("/process", methods=["POST"])
def process():
    if pythoncom:
        try:
            pythoncom.CoInitialize()
        except Exception:
            pass
            
    try:
        if 'files' not in request.files:
            return jsonify({"error": "No files uploaded."}), 400

        custom_key = request.form.get("api_key", "").strip() or request.headers.get("X-Groq-Api-Key", "").strip()
        req_api_key = custom_key if custom_key else GROQ_API_KEY
        pref_qual_text = request.form.get("preferred_qualification", "").strip()
        pref_qual_list = []
        if pref_qual_text:
            pref_qual_list = summarize_preferred_qualification(pref_qual_text, req_api_key)

        files = request.files.getlist('files')
        if not files or all(f.filename == '' for f in files):
            return jsonify({"error": "No selected files."}), 400
            
        processed_files = []
        errors = []
        
        for file in files:
            if file.filename:
                filename = secure_filename(file.filename)
                file_bytes = file.read()
                
                try:
                    payload = extract_text_from_bytes(file_bytes, filename)
                    if not payload.strip():
                        errors.append(f"{filename}: Extracted text is empty.")
                        continue
                    
                    # Fetch JSON from LLM and recursively clean any Pivot references
                    raw_data = call_groq(payload, api_key=req_api_key)
                    data = clean_data_recursively(raw_data)
                    
                    # Self-healing location extraction fallback
                    loc = data.get("location", "").strip()
                    if not loc or loc.lower() in ["none", "null", "tbd", "any", ""]:
                        match = re.search(r'(?:location|place of work)\s*[:\-–—]?\s*([a-zA-Z\s]+)', payload, re.IGNORECASE)
                        if match:
                            data["location"] = match.group(1).strip().split('\n')[0].strip()
                        else:
                            data["location"] = "Mumbai"

                    data["preferred_qualifications"] = pref_qual_list
                    safe_title = re.sub(r"[^\w\s-]", "", data.get("job_title", "Role")).replace(" ", "_")
                    outname_docx = f"{safe_title}_Role_Playbook_{uuid.uuid4().hex[:4]}.docx"

                    outpath_docx = os.path.join(app.config['OUTPUT_FOLDER'], outname_docx)
                    
                    build_docx(data, outpath_docx)
                    
                    import docx2pdf
                    outname_pdf = outname_docx.replace(".docx", ".pdf")
                    outpath_pdf = outpath_docx.replace(".docx", ".pdf")
                    docx2pdf.convert(outpath_docx, outpath_pdf)
                    
                    processed_files.append((outname_pdf, outpath_pdf))
                except Exception as e:
                    import traceback
                    print(f"Error processing {filename}: {e}")
                    traceback.print_exc()
                    errors.append(f"{filename}: {str(e)}")

        if not processed_files:
            return jsonify({"error": "Failed to process files. Errors: " + " | ".join(errors)}), 500

        if len(processed_files) == 1:
            return send_file(processed_files[0][1], as_attachment=True, download_name=processed_files[0][0])
        
        zip_path = os.path.join(app.config['OUTPUT_FOLDER'], "ABLBL_Role_Playbooks.zip")
        with zipfile.ZipFile(zip_path, 'w') as zipf:
            for fname, fpath in processed_files:
                zipf.write(fpath, arcname=fname)
                
        return send_file(zip_path, as_attachment=True, download_name="ABLBL_Role_Playbooks.zip")
    finally:
        if pythoncom:
            try:
                pythoncom.CoUninitialize()
            except Exception:
                pass

if __name__ == "__main__":
    app.run(debug=True, port=5000)
